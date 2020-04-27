from fastapi import FastAPI
from starlette.staticfiles import StaticFiles
from starlette.templating import Jinja2Templates
from starlette.requests import Request
from starlette.responses import JSONResponse
from service.handler import ServiceHandler
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from io import BytesIO

import json

handler = ServiceHandler()
api_core = FastAPI(openapi_prefix="/")
api_core.mount("/theme", StaticFiles(directory="webapp/theme"), name="static")
templates = Jinja2Templates(directory="webapp/html")


@api_core.get("/")
async def root(request: Request):
    index = dict(request=request)
    index['needs_review'] = await handler.data_svc.status_grouper("needs_review")
    index['queue'] = await handler.data_svc.status_grouper("queue")
    index['in_review'] = await handler.data_svc.status_grouper("in_review")
    index['completed'] = await handler.data_svc.status_grouper("completed")
    return templates.TemplateResponse("index.html", index)


@api_core.get("/about")
async def about(request: Request):
    return templates.TemplateResponse("about.html", {"request": request})


@api_core.post("/rest")
async def rest_api_post(request: Request):
    data = dict(await request.json())
    index = data.pop('index')
    options = dict(
        false_positive=lambda d: handler.rest_svc.false_positive(criteria=d),
        true_positive=lambda d: handler.rest_svc.true_positive(criteria=d),
        false_negative=lambda d: handler.rest_svc.false_negative(criteria=d),
        insert_report=lambda d: handler.rest_svc.insert_report(criteria=d),
        insert_csv=lambda d: handler.rest_svc.insert_csv(criteria=d),
        sentence_context=lambda d: handler.rest_svc.sentence_context(criteria=d),
        confirmed_sentences=lambda d: handler.rest_svc.confirmed_sentences(criteria=d),
    )
    result = await options[index](data)
    return result


@api_core.put("/rest")
async def report_submit_put(request: Request):
    data = dict(await request.json())
    index = data.pop('index')
    options = dict(
        set_status=lambda d: handler.rest_svc.set_status(criteria=d),
        missing_technique=lambda d: handler.rest_svc.missing_technique(criteria=d),
    )
    result = await options[index](data)
    return result


@api_core.delete("/rest")
async def rest_api_delete(request: Request):
    data = dict(await request.json())
    index = data.pop('index')
    options = dict(
        delete_report=lambda d: handler.rest_svc.delete_report(criteria=d),
        remove_sentences=lambda d: handler.rest_svc.remove_sentences(criteria=d),
    )
    result = await options[index](data)
    return result


@api_core.get("/edit/{title}")
async def report_view(request: Request,title: str):
    report = await handler.dao.get('reports', dict(title=title))
    sentences = await handler.data_svc.build_sentences(report[0]['uid'])
    attack_uids = await handler.dao.get('attack_uids')
    original_html = await handler.dao.get('original_html', dict(report_uid=report[0]['uid']))
    final_html = await handler.web_svc.build_final_html(original_html, sentences)
    output = dict(request=request,file=title, title=title, sentences=sentences, attack_uids=attack_uids, original_html=original_html, final_html=final_html)
    return templates.TemplateResponse('columns.html',output)


@api_core.get("/export/nav/{title}")
async def nav_export(request: Request, title: str):
    """
    Function to export confirmed sentences in layer json format
    :param request: THe request being sent to the server
    :param title: the title of the report
    :return: the layer json
    """
    # Get the report from the database
    report = await handler.dao.get('reports', dict(title=title))
    # Create the layer name and description
    report_title = report[0]['title']
    layer_name = f"{report_title}"
    enterprise_layer_description = f"Enterprise techniques used by {report_title}, ATT&CK"
    version = '1.0'
    if (version): # add version number if it exists
        enterprise_layer_description += f" v{version}"
    # Enterprise navigator layer
    enterprise_layer = dict(name= layer_name)
    enterprise_layer['description'] = enterprise_layer_description
    enterprise_layer['domain'] = "mitre-enterprise"
    enterprise_layer['version'] = "2.2"
    enterprise_layer['techniques'] = []
    enterprise_layer["gradient"] = { # white for nonused, blue for used
        "colors": ["#ffffff", "#66b1ff"],
        "minValue": 0,
        "maxValue": 1
    }
    enterprise_layer['legendItems'] = [{
        'label': f'used by {report_title}',
        'color': "#66b1ff"
    }]

    # Get confirmed techniques for the report from the database
    techniques = await handler.data_svc.get_confirmed_techniques(report[0]['uid'])

    # Append techniques to enterprise layer
    for technique in techniques:
        enterprise_layer['techniques'].append(technique)
        
    # Return the layer JSON in the response
    layer = json.dumps(enterprise_layer)
    return layer


@api_core.get("/export/pdf/{title}")
async def export_pdf(request: Request, title: str):
    """
    Function to export report in PDF format
    :param request: The title of the report information
    :return: response status of function
    """
    # Get the report
    report = await handler.dao.get('reports', dict(title=title))
    sentences = await handler.data_svc.build_sentences(report[0]['uid'])
    attack_uids = await handler.dao.get('attack_uids')

    dd = dict()
    dd['content'] = []
    dd['styles'] = dict()

    # Document MetaData Info
    # See https://pdfmake.github.io/docs/document-definition-object/document-medatadata/
    dd['info'] = dict()
    dd['info']['title'] = report[0]['title']
    dd['info']['creator'] = report[0]['url']

    table = {"body": []}
    table["body"].append(["ID", "Name", "Identified Sentence"])

    # Add the text to the document
    for sentence in sentences:
        dd['content'].append(sentence['text'])
        if sentence['hits']:
            for hit in sentence['hits']:
                # 'hits' object doesn't provide all the information we need, so we
                # do a makeshift join here to get that information from the attack_uid
                # list. This is ineffecient, and a way to improve this would be to perform
                # a join on the database side
                matching_attacks = [i for i in attack_uids if hit['attack_uid'] == i['uid']]
                for match in matching_attacks:
                    table["body"].append([match["tid"], match["name"], sentence['text']])

    # Append table to the end
    dd['content'].append({"table": table})
    return dd


@api_core.get("/export/word/{file}")
async def export_word(request: Request, file: str):
    """
            Function to generate JSON object containing sentence and
            sentence hit information needed to export report in Word format
            :param file: word file
            :param request: The report information
            :return: JSON response
            """
    # Get the report and report sentences data from the database
    report = await handler.dao.get('reports', dict(title=file))
    sentences = await handler.data_svc.build_sentences_for_export(report[0]['uid'])

    # Create data structure for the JSON object to be returned
    dd = {'styles': {"color": "ffc107"}, 'info': {}}

    # Add the default highlighting color style for found sentences

    # Add the report info to the data object
    dd['info']['title'] = report[0]['title']
    dd['info']['creator'] = report[0]['url']

    # Create the table object containing the attack technique information
    table = {"body": []}
    table_header = {"id": "ID", "name": "Name", "sentence": "Identified Sentence", "confirmed": "Confirmed"}
    table["body"].append(table_header)

    # Add the sentence and sentence hits information to the data object
    # and add the attack technique information to the table object
    dd['sentences'] = []
    # For every sentence returned get the sentence and sentence hits information
    # and add it to the data object to be returned
    for sentence in sentences:
        # Get the sentence information
        dd_sentence = {"text": sentence['text'], "found_status": sentence['found_status'], "hits": []}
        if sentence['hits']:
            # For every sentence hit, get the sentence hit information
            for hit in sentence['hits']:
                # Get the sentence hit information and add it to the sentence object
                dd_hit = {"attack_tid": hit['attack_tid'], "name": hit['name'], "text": sentence['text'],
                          "confirmed": hit['confirmed']}
                dd_sentence["hits"].append(dd_hit)
                # Get the attack technique information and add it to the table object
                table_row = {"id": hit['attack_tid'], "name": hit['name'], "sentence": sentence['text'],
                             "confirmed": hit['confirmed']}
                table["body"].append(table_row)
        # Add the sentence to the data object to be returned
        dd['sentences'].append(dd_sentence)

    # Add the table object to the date object to be returned
    dd['table'] = {}
    dd['table'] = table

    # Return the data object as a JSON object
    return JSONResponse(content=dd)


@api_core.route("/export/word/doc", '*')
async def export_word_doc(request: Request):
    """
            Function to generate the docx object containing sentence and
            sentence hit information to export report in Word
            :param request: The JSON report information
            :return: A byte array response
            """

    # Get the JSON object containing the report information
    data = dict(await request.json())

    # Create the docx Document object
    document = Document()

    # Add the title and heading to the document object
    document.core_properties.title = data['info']['title']
    document.add_heading(data['info']['title'], 1)
    document.add_paragraph('')

    # Get the sentences information from the data object
    sentences = data['sentences']
    # Add a new paragraph to the document object to contain the sentences
    paragraph = document.add_paragraph()

    # Add each sentence to the paragraph in the document
    for sentence in sentences:
        # Get the found_status attribute of the sentence
        found_status = sentence['found_status']
        if found_status == 'true':
            # if the sentence found status attribute is true then
            # add the sentence with additional newlines and a '*' and
            # highlighting to indicate that the sentence includes attack techniques
            run = paragraph.add_run("\n* " + sentence['text'] + "\n")
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            # otherwise, just add the sentence to the paragraph
            run = paragraph.add_run(sentence['text'])

    # Add a table to the document object containing the attack techniques found
    # Get the items in the table object from the data object
    items = data['table']['body']
    # Add a table to the document containing the attack technique found information
    table = document.add_table(rows=0, cols=4, style='LightGrid')
    # For each item in the table object, add a new row to the table
    for item in items:
        # Add d new row to the table
        cells = table.add_row().cells
        # If the item is the first item containing the table headers
        # then add the table header row to the table
        if item['id'] == 'ID':
            run = cells[0].paragraphs[0].add_run(item['id'])
            run.bold = True
            run = cells[1].paragraphs[0].add_run(item['name'])
            run.bold = True
            run = cells[2].paragraphs[0].add_run(item['sentence'])
            run.bold = True
            run = cells[3].paragraphs[0].add_run(item['confirmed'])
            run.bold = True
        else:
            # otherwise, add the item information to the row
            cells[0].text = item['id']
            cells[1].text = item['name']
            cells[2].text = item['sentence']
            cells[3].text = item['confirmed']

    # Save the document object to a byte stream
    stream = BytesIO()
    document.save(stream)
    stream.seek(0)

    # Read the document object from the byte stream to a byte array
    b = bytearray(stream.read())

    # Return the byte array
    return b
