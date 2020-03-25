from aiohttp_jinja2 import template, web
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from io import BytesIO
import nltk
import json


class WebAPI:

    def __init__(self, services):

        self.dao = services.get('dao')
        self.data_svc = services['data_svc']
        self.web_svc = services['web_svc']
        self.ml_svc = services['ml_svc']
        self.reg_svc = services['reg_svc']
        self.rest_svc = services['rest_svc']
        self.tokenizer_sen = nltk.data.load('tokenizers/punkt/english.pickle')

    @template('about.html')
    async def about(self, request):
        return

    @template('index.html')
    async def index(self, request):
        index = dict(needs_review=await self.data_svc.status_grouper("needs_review"))
        index['queue'] = await self.data_svc.status_grouper("queue")
        index['in_review'] = await self.data_svc.status_grouper("in_review")
        index['completed'] = await self.data_svc.status_grouper("completed")
        return index

    async def rest_api(self, request):
        """
        Function to handle rest api calls
        :param request: json data with rest request
        :return: json response
        """
        data = dict(await request.json())
        index = data.pop('index')
        options = dict(
            POST=dict(
                false_positive=lambda d: self.rest_svc.false_positive(criteria=d),
                true_positive=lambda d: self.rest_svc.true_positive(criteria=d),
                false_negative=lambda d: self.rest_svc.false_negative(criteria=d),
                set_status=lambda d: self.rest_svc.set_status(criteria=d),
                insert_report=lambda d: self.rest_svc.insert_report(criteria=d),
                insert_csv=lambda d: self.rest_svc.insert_csv(criteria=d),
                remove_sentences=lambda d: self.rest_svc.remove_sentences(criteria=d),
                delete_report=lambda d: self.rest_svc.delete_report(criteria=d),
                sentence_context=lambda d: self.rest_svc.sentence_context(criteria=d),
                confirmed_sentences=lambda d: self.rest_svc.confirmed_sentences(criteria=d),
                missing_technique=lambda d: self.rest_svc.missing_technique(criteria=d)
            ))
        output = await options[request.method][index](data)
        return web.json_response(output)

    @template('columns.html')
    async def edit(self, request):
        """
        Function to edit report
        :param request: The title of the report information
        :return: dictionary of report data
        """
        report = await self.dao.get('reports', dict(title=request.match_info.get('file')))
        sentences = await self.data_svc.build_sentences(report[0]['uid'])
        attack_uids = await self.dao.get('attack_uids')
        original_html = await self.dao.get('original_html', dict(report_uid=report[0]['uid']))
        final_html = await self.web_svc.build_final_html(original_html, sentences)
        return dict(file=request.match_info.get('file'), title=report[0]['title'], sentences=sentences, attack_uids=attack_uids, original_html=original_html, final_html=final_html)

    async def nav_export(self, request):
        """
        Function to export confirmed sentences in layer json format
        :param request: The title of the report information
        :return: the layer json
        """        
        # Get the report from the database
        report = await self.dao.get('reports', dict(title=request.match_info.get('file')))

        # Create the layer name and description
        report_title = report[0]['title']
        layer_name = f"{report_title}"
        enterprise_layer_description = f"Enterprise techniques used by {report_title}, ATT&CK"
        version = '1.0'
        if (version): # add version number if it exists
            enterprise_layer_description += f" v{version}"

        # Enterprise navigator layer
        enterprise_layer = {}
        enterprise_layer['name'] = layer_name
        enterprise_layer['description'] = enterprise_layer_description
        enterprise_layer['domain'] = "mitre-enterprise"
        enterprise_layer['version'] = "2.2"
        enterprise_layer['techniques'] = []
        enterprise_layer["gradient"] = { # white for nonused, blue for used
		    "colors": ["#ffffff", "#66b1ff"],
		    "minValue": 0,
    		"maxValue": 1
	    }
        enterprise_layer['legendItems'] = [{
            'label': f'used by {report_title}',
            'color': "#66b1ff"
        }]

        # Get confirmed techniques for the report from the database
        techniques = await self.data_svc.get_confirmed_techniques(report[0]['uid'])

        # Append techniques to enterprise layer
        for technique in techniques:
            enterprise_layer['techniques'].append(technique)
            
        # Return the layer JSON in the response
        layer = json.dumps(enterprise_layer)
        return web.json_response(layer)

    async def word_export(self, request):
        """
        Function to generate JSON object containing sentence and
        sentence hit information needed to export report in Word format
        :param request: The report information
        :return: JSON response
        """
        # Get the report and report sentences data from the database
        report = await self.dao.get('reports', dict(title=request.match_info.get('file')))
        sentences = await self.data_svc.build_sentences_for_export(report[0]['uid'])
        
        # Create data structure for the JSON object to be returned
        dd = {}
        
        # Add the default highlighting color style for found sentences
        dd['styles'] = {"color":"ffc107"}

        # Add the report info to the data object
        dd['info'] = {}
        dd['info']['title'] = report[0]['title']
        dd['info']['creator'] = report[0]['url']
        
        # Create the table object containing the attack technique information
        table = {"body": []}
        table_header = {"id":"ID", "name":"Name", "sentence":"Identified Sentence", "confirmed":"Confirmed"}
        table["body"].append(table_header)

        # Add the sentence and sentence hits information to the data object
        # and add the attack technique information to the table object
        dd['sentences'] = []
        # For every sentence returned get the sentence and sentence hits information
        # and add it to the data object to be returned
        for sentence in sentences:
            # Get the sentence information
            dd_sentence = {"text":sentence['text'], "found_status":sentence['found_status'], "hits": []}
            if sentence['hits']:
                # For every sentence hit, get the sentence hit information
                for hit in sentence['hits']:
                    # Get the sentence hit information and add it to the sentence object
                    dd_hit = {"attack_tid":hit['attack_tid'], "name":hit['name'], "text":sentence['text'], "confirmed":hit['confirmed']}
                    dd_sentence["hits"].append(dd_hit)
                    # Get the attack technique information and add it to the table object
                    table_row = {"id":hit['attack_tid'], "name":hit['name'], "sentence":sentence['text'], "confirmed":hit['confirmed']}
                    table["body"].append(table_row)
            # Add the sentence to the data object to be returned
            dd['sentences'].append(dd_sentence)

        # Add the table object to the date object to be returned
        dd['table'] = {}
        dd['table'] = table

        # Return the data object as a JSON object
        return web.json_response(dd)

    async def export_to_word(self, request):
        """
        Function to generate the docx object containing sentence and
        sentence hit information to export report in Word
        :param request: The JSON report information
        :return: A byte array response
        """

        # Get the JSON object containing the report information
        data = dict(await request.json())
        
        # Create the docx Document object
        document = Document()

        # Add the title and heading to the document object
        document.core_properties.title = data['info']['title']
        document.add_heading(data['info']['title'], 1)
        document.add_paragraph('')

        # Get the sentences information from the data object
        sentences = data['sentences']
        # Add a new paragraph to the document object to contain the sentences
        paragraph = document.add_paragraph()
        
        # Add each sentence to the paragraph in the document
        for sentence in sentences:
            # Get the found_status attribute of the sentence
            found_status = sentence['found_status']
            if found_status == 'true':
                # if the sentence found status attribute is true then
                # add the sentence with additional newlines and a '*' and
                # highlighting to indicate that the sentence includes attack techniques
                run = paragraph.add_run("\n* " + sentence['text'] + "\n")
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                # otherwise, just add the sentence to the paragraph
                run = paragraph.add_run(sentence['text'])

        # Add a table to the document object containing the attack techniques found
        # Get the items in the table object from the data object
        items = data['table']['body']
        # Add a table to the document containing the attack technique found information
        table = document.add_table(rows=0, cols=4, style='LightGrid')
        # For each item in the table object, add a new row to the table
        for item in items:
            # Add d new row to the table
            cells = table.add_row().cells
            # If the item is the first item containing the table headers
            # then add the table header row to the table
            if item['id'] == 'ID':
                run = cells[0].paragraphs[0].add_run(item['id'])
                run.bold = True
                run = cells[1].paragraphs[0].add_run(item['name'])
                run.bold = True
                run = cells[2].paragraphs[0].add_run(item['sentence'])
                run.bold = True
                run = cells[3].paragraphs[0].add_run(item['confirmed'])
                run.bold = True
            else:
                # otherwise, add the item information to the row
                cells[0].text = item['id']
                cells[1].text = item['name']
                cells[2].text = item['sentence']
                cells[3].text = item['confirmed']
        
        # Save the document object to a byte stream
        stream = BytesIO()
        document.save(stream)
        stream.seek(0)

        # Read the document object from the byte stream to a byte array
        b = bytearray(stream.read())
        
        # Return the byte array
        return web.Response(body=b)
        
    async def pdf_export(self, request):
        """
        Function to export report in PDF format
        :param request: The report information
        :return: JSON response
        """
        # Get the reports
        report = await self.dao.get('reports', dict(title=request.match_info.get('file')))
        sentences = await self.data_svc.build_sentences_for_export(report[0]['uid'])

        dd = dict()
        dd['content'] = []
        dd['styles'] = {"color":"ffc107"}

        # Document MetaData Info
        # See https://pdfmake.github.io/docs/document-definition-object/document-medatadata/
        dd['info'] = dict()
        dd['info']['title'] = report[0]['title']
        dd['info']['creator'] = report[0]['url']

        table = {"body": []}
        table["body"].append(["ID", "Name", "Identified Sentence", "Confirmed"])

        # Add the text to the document
        for sentence in sentences:
            if sentence['found_status'] == "true":
                text = {"text": sentence['text'], "style": "color"}
            else:
                text = {"text": sentence['text']}
            dd['content'].append(text)
            if sentence['hits']:
                for hit in sentence['hits']:
                    table["body"].append([hit['attack_tid'], hit['name'], sentence['text'], hit['confirmed']])

        # Append table to the end
        dd['content'].append({"table": table})
        return web.json_response(dd)

    async def rebuild_ml(self, request):
        """
        This is a new api function to force a rebuild of the ML models. This is intended to be kicked off in the background at some point
        :param request: uh, nothing?
        :return: status of rebuild
        """
        # get techniques from database
        tech_data = await self.dao.get('attack_uids')
        techniques = {}
        for row in tech_data:
            # skip software for now
            if 'tool' in row['tid'] or 'malware' in row['tid']:
                continue
            else:
                # query for true positives
                true_pos = await self.dao.get('true_positives', dict(uid=row['uid']))
                tp = []
                for t in true_pos:
                    tp.append(t['true_positive'])
                # query for false negatives and false positives
                false_neg = await self.dao.get('false_negatives', dict(uid=row['uid']))
                false_positives = await self.dao.get('false_positives', dict(uid=row['uid']))
                for f in false_neg:
                    tp.append(f['false_negative'])
                fp = []
                for fps in false_positives:
                    fp.append(fps['false_positive'])

                techniques[row['uid']] = {'id': row['tid'], 'name': row['name'], 'similar_words': [],
                                          'example_uses': tp, 'false_positives': fp}

        # query for true negatives
        true_negatives = []
        true_negs = await self.dao.get('true_negatives')
        for i in true_negs:
            true_negatives.append(i['sentence'])
        list_of_legacy, list_of_techs = await self.data_svc.ml_reg_split(techniques)
        self.ml_svc.build_pickle_file(self, list_of_techs, techniques, true_negatives, force=True)

        return {'text': 'ML Rebuilt!'}


